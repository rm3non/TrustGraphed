"""
TrustGraphed Evaluation Routes
"""

from flask import Blueprint, request, jsonify
import sys
import os
import tempfile
import fitz  # PyMuPDF
import docx

# Add the backend directory to the Python path
backend_dir = os.path.dirname(os.path.abspath(__file__))
backend_parent = os.path.dirname(backend_dir)
if backend_parent not in sys.path:
    sys.path.insert(0, backend_parent)

# Import utils modules
from utils.sdg import SourceDataGrappler
from utils.aie import AssertionIntegrityEngine
from utils.cce import ConfidenceComputationEngine
from utils.zfp import ZeroFabricationProtocol
from utils.score_engine import TrustScoreEngine
from utils.certificate import CertificateGenerator

evaluate_bp = Blueprint('evaluate', __name__)

def extract_text_from_file(file):
    """Extract text content from uploaded file with comprehensive error handling."""
    filename = file.filename.lower() if file.filename else ""
    print(f"Processing file: {filename}")

    if not filename:
        raise ValueError("No filename provided")

    try:
        if filename.endswith(('.txt', '.md')):
            # Handle text files
            content = file.read().decode('utf-8', errors='replace')
            if not content.strip():
                raise ValueError("Text file appears to be empty")
            return content

        elif filename.endswith('.pdf'):
            # Handle PDF files using PyMuPDF (fitz)
            content = ""
            try:
                # Read file content into bytes
                file_bytes = file.read()
                if not file_bytes:
                    raise ValueError("PDF file appears to be empty")

                # Open PDF document from bytes
                pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")

                if pdf_doc.page_count == 0:
                    raise ValueError("PDF has no pages")

                # Extract text from all pages
                for page_num in range(pdf_doc.page_count):
                    page = pdf_doc[page_num]
                    page_text = page.get_text()
                    if page_text.strip():  # Only add non-empty pages
                        content += page_text + "\n"

                pdf_doc.close()

                if not content.strip():
                    raise ValueError("No readable text found in PDF")

                return content

            except Exception as pdf_error:
                raise ValueError(f"PDF processing failed: {str(pdf_error)}")

        elif filename.endswith('.docx'):
            # Handle DOCX files
            try:
                doc = docx.Document(file)
                content = ""

                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content += paragraph.text + "\n"

                # Extract text from tables if any
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                content += cell.text + " "
                    content += "\n"

                if not content.strip():
                    raise ValueError("No readable text found in DOCX file")

                return content

            except Exception as docx_error:
                raise ValueError(f"DOCX processing failed: {str(docx_error)}")

        elif filename.endswith('.doc'):
            # Legacy DOC files - basic text extraction attempt
            try:
                content = file.read().decode('utf-8', errors='replace')
                # Remove common binary artifacts
                content = ''.join(char for char in content if char.isprintable() or char.isspace())

                if not content.strip() or len(content.strip()) < 10:
                    raise ValueError("Unable to extract readable text from DOC file")

                return content

            except Exception as doc_error:
                raise ValueError(f"DOC processing failed: {str(doc_error)}")

        else:
            supported_types = ['.txt', '.md', '.pdf', '.docx', '.doc']
            raise ValueError(f"Unsupported file type. Supported formats: {', '.join(supported_types)}")

    except ValueError:
        # Re-raise ValueError as-is
        raise
    except Exception as e:
        # Catch any other unexpected errors
        raise ValueError(f"Unexpected error processing file '{filename}': {str(e)}")

evaluate_bp = Blueprint('evaluate', __name__)

@evaluate_bp.route('/evaluate', methods=['POST'])
def evaluate_content():
    """
    Main evaluation endpoint - processes content through all 6 TrustGraphed modules.
    """
    try:
        # Handle both file uploads and text input
        content = ""

        if 'file' in request.files:
            # File upload mode
            uploaded_file = request.files['file']

            if uploaded_file.filename == '' or uploaded_file.filename is None:
                return jsonify({"error": "No file selected"}), 400

            # Validate file size (10MB limit)
            file_size = len(uploaded_file.read())
            uploaded_file.seek(0)  # Reset file pointer

            if file_size == 0:
                return jsonify({"error": "File is empty"}), 400

            if file_size > 10 * 1024 * 1024:  # 10MB limit
                return jsonify({"error": "File size exceeds 10MB limit"}), 400

            # Extract text from file
            try:
                content = extract_text_from_file(uploaded_file)

                # Log successful extraction for debugging
                print(f"Successfully extracted {len(content)} characters from {uploaded_file.filename}")

            except ValueError as e:
                error_msg = str(e) if str(e) else "Unknown file processing error"
                print(f"ValueError processing file '{uploaded_file.filename}': {error_msg}")
                return jsonify({
                    "error": f"File processing error: {error_msg}",
                    "filename": uploaded_file.filename,
                    "status": "failed"
                }), 400
            except Exception as e:
                error_msg = str(e) if str(e) else "Unknown unexpected error"
                print(f"Exception processing file '{uploaded_file.filename}': {error_msg}")
                print(f"Exception type: {type(e).__name__}")
                return jsonify({
                    "error": f"Unexpected file processing error: {error_msg}",
                    "filename": uploaded_file.filename,
                    "status": "failed"
                }), 500

        elif request.is_json:
            # Text input mode
            data = request.get_json()
            if not data or 'content' not in data:
                return jsonify({"error": "Missing 'content' field in request body"}), 400
            content = data['content']
        else:
            return jsonify({"error": "Please provide content via JSON or file upload"}), 400

        if not content or len(content.strip()) < 10:
            return jsonify({"error": "Content must be at least 10 characters long"}), 400

        # Initialize all modules
        sdg = SourceDataGrappler()
        aie = AssertionIntegrityEngine()
        cce = ConfidenceComputationEngine()
        zfp = ZeroFabricationProtocol()
        score_engine = TrustScoreEngine()
        certificate_gen = CertificateGenerator()

        # Process through pipeline
        results = {}

        # Step 1: Extract assertions and citations
        sdg_result = sdg.process(content)
        results['sdg_result'] = sdg_result

        # Step 2: Check assertion integrity
        assertions = sdg_result.get('assertions', [])
        aie_result = aie.process(content, assertions)
        results['aie_result'] = aie_result

        # Step 3: Compute confidence scores
        citations = sdg_result.get('citations', [])
        cce_result = cce.process(content, assertions, citations)
        results['cce_result'] = cce_result

        # Step 4: Check for fabrication
        zfp_result = zfp.process(content)
        results['zfp_result'] = zfp_result

        # Step 5: Generate final trust score
        score_result = score_engine.process(results)
        results['score_result'] = score_result

        # Step 6: Generate certificate
        cert_result = certificate_gen.process(content, score_result)
        results['certificate_result'] = cert_result

        # Build response
        response = {
            "status": "success",
            "content_length": len(content),
            "trust_evaluation": {
                "trust_score": score_result['trust_score'],
                "trust_level": score_result['trust_level'],
                "component_scores": score_result['component_scores'],
                "insights": score_result['insights'],
                "detailed_explanation": score_result.get('detailed_explanation', {})
            },
            "certificate_id": cert_result['certificate_id'],
            "module_results": {
                "source_data_grappler": {
                    "assertions_found": sdg_result['assertions_count'],
                    "citations_found": sdg_result['citations_count'],
                    "extraction_confidence": sdg_result['extraction_confidence']
                },
                "assertion_integrity": {
                    "integrity_score": aie_result['integrity_score'],
                    "issues_found": aie_result['issues_found']
                },
                "confidence_computation": {
                    "overall_confidence": cce_result['overall_confidence'],
                    "high_confidence_assertions": cce_result['high_confidence_count']
                },
                "zero_fabrication": {
                    "authenticity_score": zfp_result['authenticity_score'],
                    "fabrication_risk": zfp_result['fabrication_risk'],
                    "flags_detected": zfp_result['total_flags']
                }
            },
            "certificate": cert_result['certificate'],
            "readable_summary": cert_result['readable_summary']
        }

        return jsonify(response), 200

    except Exception as e:
        error_message = str(e) if str(e) else "Unknown processing error occurred"
        print(f"Error during evaluation: {error_message}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'error': f'Processing error: {error_message}',
            'status': 'error',
            'details': 'Please check file format and try again'
        }), 500

@evaluate_bp.route('/evaluate/health', methods=['GET'])
def evaluate_health():
    """Health check for evaluation service."""
    return jsonify({
        "status": "healthy",
        "service": "TrustGraphed Evaluation Pipeline",
        "modules": [
            "Source Data Grappler",
            "Assertion Integrity Engine", 
            "Confidence Computation Engine",
            "Zero-Fabrication Protocol",
            "TrustScore Engine",
            "Certificate Generator"
        ]
    })

@evaluate_bp.route('/evaluate/test-file', methods=['POST'])
def test_file_processing():
    """Test endpoint for file processing without full evaluation."""
    try:
        print("Test file processing endpoint called")
        print(f"Request files: {list(request.files.keys())}")

        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400

        uploaded_file = request.files['file']
        print(f"Uploaded file: {uploaded_file.filename}")

        if not uploaded_file.filename or uploaded_file.filename == '':
            return jsonify({"error": "No file selected"}), 400

        # Extract text from file
        try:
            content = extract_text_from_file(uploaded_file)

            return jsonify({
                "status": "success",
                "filename": uploaded_file.filename,
                "content_length": len(content),
                "content_preview": content[:200] + "..." if len(content) > 200 else content,
                "message": "File processed successfully"
            }), 200

        except ValueError as e:
            error_msg = str(e) if str(e) else "Unknown ValueError"
            print(f"ValueError in test endpoint: {error_msg}")
            print(f"File details - Name: {uploaded_file.filename}, Size: {len(uploaded_file.read()) if uploaded_file else 'N/A'}")
            uploaded_file.seek(0) if uploaded_file else None
            return jsonify({
                "error": error_msg,
                "filename": uploaded_file.filename,
                "debug_info": {
                    "error_type": "ValueError",
                    "file_size": len(uploaded_file.read()) if uploaded_file else 0
                }
            }), 400
        except Exception as e:
            error_msg = str(e) if str(e) else "Unknown Exception"
            print(f"Exception in test endpoint: {error_msg}")
            print(f"Exception type: {type(e).__name__}")
            return jsonify({
                "error": error_msg,
                "filename": uploaded_file.filename if uploaded_file else "unknown",
                "debug_info": {
                    "error_type": type(e).__name__,
                    "file_size": len(uploaded_file.read()) if uploaded_file else 0
                }
            }), 500

    except Exception as e:
        error_msg = str(e) if str(e) else "Unknown test failure"
        print(f"Test endpoint exception: {error_msg}")
        return jsonify({"error": f"Test failed: {error_msg}"}), 500